import sys
import os
os.environ['PYTHONIOENCODING'] = 'utf-8'
from docx import Document
from docx.oxml.ns import qn

def rgb_to_hex(color):
    if color is None:
        return None
    if color.rgb:
        return str(color.rgb)
    return None

def analyze_docx(path):
    doc = Document(path)
    print("=" * 70)
    print(f"ANALYZE: {path}")
    print("=" * 70)

    print("\n[PAGE SETUP]")
    for section in doc.sections:
        print(f"  Page size: {section.page_width.pt:.1f} pt x {section.page_height.pt:.1f} pt")
        print(f"  Margins: L={section.left_margin.pt:.1f}, R={section.right_margin.pt:.1f}, T={section.top_margin.pt:.1f}, B={section.bottom_margin.pt:.1f}")

    print("\n[STYLES IN DOCUMENT]")
    styles_seen = {}
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "None"
        if style_name not in styles_seen:
            styles_seen[style_name] = {"count": 0, "examples": []}
        styles_seen[style_name]["count"] += 1
        if len(styles_seen[style_name]["examples"]) < 2:
            snippet = para.text[:80].replace("\n", " ")
            styles_seen[style_name]["examples"].append(snippet)

    for name, info in sorted(styles_seen.items(), key=lambda x: -x[1]["count"]):
        print(f"\n  Style: '{name}' -- used {info['count']} times")
        for ex in info["examples"]:
            print(f"      > {ex}")

    print("\n[PARAGRAPH FORMAT DETAILS (first 50 non-empty paragraphs)]")
    shown = 0
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue
        if shown >= 50:
            break
        shown += 1
        pf = para.paragraph_format
        alignment = str(pf.alignment)
        space_after = pf.space_after.pt if pf.space_after else None
        space_before = pf.space_before.pt if pf.space_before else None
        left_indent = pf.left_indent.pt if pf.left_indent else 0
        first_indent = pf.first_line_indent.pt if pf.first_line_indent else 0
        style = para.style.name if para.style else "None"

        font_info = []
        if para.runs:
            for run in para.runs[:3]:
                if run.text.strip():
                    f = run.font
                    font_info.append(
                        f"{f.name}|{f.size.pt if f.size else '?'}pt|bold={f.bold}|color={rgb_to_hex(f.color)}"
                    )

        print(f"  [{i}] Style={style:<25} | Align={alignment:<12} | SpaceBefore={space_before} | SpaceAfter={space_after} | LeftIndent={left_indent:.1f} | FirstIndent={first_indent:.1f}")
        print(f"       Text: {para.text[:100]}")
        if font_info:
            print(f"       Font: {'  ;  '.join(font_info)}")

    print("\n[TABLES]")
    for i, table in enumerate(doc.tables):
        print(f"\n  Table {i+1}: {len(table.rows)} rows x {len(table.columns)} cols")
        for r_idx, row in enumerate(table.rows[:3]):
            cells = []
            for cell in row.cells:
                text = cell.text[:30].replace("\n", " ")
                shading = cell._tc.get_or_add_tcPr()
                shd = shading.find(qn('w:shd'))
                bg = shd.get(qn('w:fill')) if shd is not None else None
                cells.append(f"[{text}](bg={bg})")
            print(f"    Row {r_idx}: {' | '.join(cells)}")
        if len(table.rows) > 3:
            print(f"    ... ({len(table.rows)-3} more rows)")

    print("\n[ALL FONTS USED]")
    all_fonts = {}
    for para in doc.paragraphs:
        for run in para.runs:
            fname = run.font.name
            fsize = run.font.size.pt if run.font.size else None
            fbold = run.font.bold
            key = (fname, fsize, fbold)
            all_fonts[key] = all_fonts.get(key, 0) + 1

    for (fname, fsize, fbold), count in sorted(all_fonts.items(), key=lambda x: -x[1]):
        print(f"  {count:5d} chars: Font={fname}, Size={fsize}pt, Bold={fbold}")

    print("\n" + "=" * 70)

if __name__ == "__main__":
    analyze_docx(sys.argv[1])
